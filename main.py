import docx

doc = docx.Document('source_document_for_multiple_needs__roles_with_positions_amount.docx')


doc.paragraphs[1].runs[0].style = 'Intense Emphasis'
doc.paragraphs[1].runs[4].underline = True

doc.save('restyled2.docx')